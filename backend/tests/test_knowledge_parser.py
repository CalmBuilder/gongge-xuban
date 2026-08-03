"""
@Time       : 2026/07/27 14:35
@Author     : zhanglp8181
@File       : test_knowledge_parser.py
@CallChain  : pytest → knowledge.parser.extract_text → DOCX/PDF 结构文本
@Description: 验证知识文档解析保留标题和分页结构。
"""

from io import BytesIO

from docx import Document

from app.knowledge.parser import _docx_heading_level, extract_text


def test_docx_parser_preserves_heading_hierarchy() -> None:
    """DOCX 标题样式应转换为 Markdown 层级供章节树继续消费。"""

    document = Document()
    document.add_heading("员工手册", level=1)
    document.add_heading("请假制度", level=2)
    document.add_paragraph("正式员工可按制度申请年假。")
    buffer = BytesIO()
    document.save(buffer)

    text, file_type = extract_text("employee-handbook.docx", buffer.getvalue())

    assert file_type == "docx"
    assert "# 员工手册" in text
    assert "## 请假制度" in text
    assert "正式员工可按制度申请年假。" in text


def test_docx_heading_level_supports_chinese_and_english_styles() -> None:
    """标题层级识别同时兼容 Word 的中英文标准样式名。"""

    assert _docx_heading_level("Heading 3") == 3
    assert _docx_heading_level("标题 2") == 2
    assert _docx_heading_level("正文") is None
