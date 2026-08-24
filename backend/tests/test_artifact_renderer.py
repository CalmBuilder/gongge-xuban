"""
@Time       : 2026/08/13 23:35
@Author     : zhanglp8181
@File       : test_artifact_renderer.py
@CallChain  : pytest → verified markdown renderer → python-docx/openpyxl round trip
@Description: 验证 DOCX/XLSX 真实产物可重读且危险表格公式只作为文本保存。
"""

import csv
from datetime import timedelta
from io import BytesIO, StringIO
from pathlib import Path

import pytest
from docx import Document
from openpyxl import load_workbook
from sqlalchemy.pool import StaticPool
from sqlmodel import Session, SQLModel, create_engine

from app.db.models import ExecutionResult, SopInstance, SopNodeExecution, User, utc_now
from app.dynamic_tasks.artifact_renderer import (
    CSV_MIME,
    DOCX_MIME,
    XLSX_MIME,
    ArtifactRenderError,
    ArtifactRendererService,
    render_verified_markdown,
)
from app.dynamic_tasks.artifacts import ArtifactService


def _session() -> Session:
    """创建覆盖RendererJob与Artifact的隔离SQLite会话。"""

    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    SQLModel.metadata.create_all(engine)
    return Session(engine)


def _facts(db: Session) -> tuple[SopInstance, SopNodeExecution, ExecutionResult]:
    """建立一个已验证结果及仍可登记required Artifact的answer节点。"""

    instance = SopInstance(
        id="render-execution",
        tenant_id="tenant-render",
        session_id="session-render",
        initiator_user_id="user-render",
        kind="sop",
        skill_id="renderer_sop",
        skill_version_id="renderer_sop_v1",
        skill_version="1.0.0",
        definition_checksum="b" * 64,
        status="running",
        active_slot_key="render-execution",
    )
    node = SopNodeExecution(
        id="render-node",
        tenant_id=instance.tenant_id,
        instance_id=instance.id,
        node_id="answer",
        step_key="answer",
        status="running",
    )
    result = ExecutionResult(
        id="render-result",
        tenant_id=instance.tenant_id,
        execution_id=instance.id,
        status="verified",
        result_json={"markdown": "# 报告"},
        verification_json={"passed": True},
        checksum="a" * 64,
    )
    user = User(
        id=instance.initiator_user_id,
        tenant_id=instance.tenant_id,
        username="renderer-user",
        password_hash="test",
        membership_status="active",
    )
    db.add_all([user, instance, node, result])
    db.commit()
    return instance, node, result


def test_docx_renderer_round_trips_verified_report() -> None:
    """DOCX 必须能由标准库重新打开并保留关键标题与事实。"""

    data = render_verified_markdown("# 续约风险报告\n合同要求提前60天通知。", DOCX_MIME)
    document = Document(BytesIO(data))
    assert [paragraph.text for paragraph in document.paragraphs] == [
        "续约风险报告",
        "合同要求提前60天通知。",
    ]


def test_xlsx_renderer_escapes_formula_and_dde_prefixes() -> None:
    """XLSX 内模型文本即使以公式/DDE前缀开头也不得成为可执行公式。"""

    data = render_verified_markdown("# 报告\n=HYPERLINK(\"bad\")\n\t@SUM(A1)", XLSX_MIME)
    workbook = load_workbook(BytesIO(data), data_only=False)
    sheet = workbook["报告"]
    assert sheet["B2"].value == "# 报告"
    assert sheet["B3"].value == "'=HYPERLINK(\"bad\")"
    assert sheet["B4"].value == "'\t@SUM(A1)"
    assert sheet["B3"].data_type == "s"


def test_csv_renderer_round_trips_and_escapes_formula_prefixes() -> None:
    """CSV 必须可由标准库重读，并把前导空白后的公式与DDE内容保持为纯文本。"""

    data = render_verified_markdown("# 报告\n=HYPERLINK(\"bad\")\n\t@SUM(A1)", CSV_MIME)
    rows = list(csv.reader(StringIO(data.decode("utf-8-sig"))))

    assert rows == [
        ["序号", "内容"],
        ["1", "# 报告"],
        ["2", "'=HYPERLINK(\"bad\")"],
        ["3", "'\t@SUM(A1)"],
    ]


def test_renderer_job_is_idempotent_and_old_worker_is_fenced(tmp_path: Path) -> None:
    """正反向证明RendererJob幂等发布DOCX，错误owner不能生成或覆盖Artifact。"""

    with _session() as db:
        instance, node, result = _facts(db)
        service = ArtifactRendererService(
            db,
            artifact_service=ArtifactService(db, storage_root=tmp_path),
        )
        job, created = service.ensure_job(
            instance=instance,
            result_id=result.id,
            result_checksum=result.checksum,
            source_node=node,
            artifact_key="report",
            filename="report.docx",
            mime_type=DOCX_MIME,
            required=True,
        )
        service.claim(job, worker_id="renderer-a")
        with pytest.raises(ArtifactRenderError, match="ARTIFACT_RENDER_JOB_FENCED"):
            service.render_and_publish(
                job,
                markdown="# 报告",
                worker_id="renderer-b",
                fencing_token=job.fencing_token,
                input_snapshot_ids=(),
            )
        artifact = service.render_and_publish(
            job,
            markdown="# 报告",
            worker_id="renderer-a",
            fencing_token=job.fencing_token,
            input_snapshot_ids=(),
        )
        replay, replay_created = service.ensure_job(
            instance=instance,
            result_id=result.id,
            result_checksum=result.checksum,
            source_node=node,
            artifact_key="report",
            filename="report.docx",
            mime_type=DOCX_MIME,
            required=True,
        )
        db.commit()
        replay_status = replay.status
        replay_artifact_id = replay.artifact_id
        replay_staged_checksum = replay.staged_checksum
        artifact_id = artifact.id
        artifact_checksum = artifact.content_checksum

    assert created is True
    assert replay_created is False
    assert replay_status == "ready"
    assert replay_artifact_id == artifact_id
    assert replay_staged_checksum == artifact_checksum


def test_expired_renderer_job_is_recovered_and_old_worker_cannot_publish(tmp_path: Path) -> None:
    """过期claimed作业必须换fencing重领，旧worker不能在恢复后迟到发布。"""

    with _session() as db:
        instance, node, result = _facts(db)
        service = ArtifactRendererService(
            db,
            artifact_service=ArtifactService(db, storage_root=tmp_path),
        )
        job, _ = service.ensure_job(
            instance=instance,
            result_id=result.id,
            result_checksum=result.checksum,
            source_node=node,
            artifact_key="recover-report",
            filename="recover.docx",
            mime_type=DOCX_MIME,
            required=True,
        )
        service.claim(job, worker_id="renderer-old")
        old_token = job.fencing_token
        job.lease_expires_at = utc_now() - timedelta(seconds=1)
        db.add(job)
        db.commit()

        assert service.requeue_expired() == 1
        db.commit()
        db.refresh(job)
        service.claim(job, worker_id="renderer-new")
        db.commit()

        with pytest.raises(ArtifactRenderError, match="ARTIFACT_RENDER_JOB_FENCED"):
            service.render_and_publish(
                job,
                markdown="# 报告",
                worker_id="renderer-old",
                fencing_token=old_token,
                input_snapshot_ids=(),
            )
        artifact = service.resume_job(job, worker_id="renderer-new")
        db.commit()

        assert artifact.status == "ready"
        _, content = ArtifactService(db, storage_root=tmp_path).resolve(
            artifact.id,
            tenant_id=instance.tenant_id,
            actor_user_id=instance.initiator_user_id,
        )
        assert Document(BytesIO(content)).paragraphs[0].text == "报告"


def test_renderer_reuses_identical_orphan_blob_after_database_rollback(tmp_path: Path) -> None:
    """模拟文件已原子发布但DB回滚，重放必须复用同一内容地址而不是永久失败。"""

    with _session() as db:
        instance, node, result = _facts(db)
        artifacts = ArtifactService(db, storage_root=tmp_path)
        service = ArtifactRendererService(db, artifact_service=artifacts)
        job, _ = service.ensure_job(
            instance=instance,
            result_id=result.id,
            result_checksum=result.checksum,
            source_node=node,
            artifact_key="orphan-report",
            filename="orphan.docx",
            mime_type=DOCX_MIME,
            required=True,
        )
        db.commit()
        data = render_verified_markdown("# 报告", DOCX_MIME)
        artifacts.register(
            instance=instance,
            source_node=node,
            artifact_key=job.artifact_key,
            filename=job.filename,
            mime_type=job.mime_type,
            data=data,
            artifact_id=job.artifact_id,
        )
        db.rollback()

        job = db.get(type(job), job.id)
        assert job is not None
        service.claim(job, worker_id="renderer-recovery")
        db.commit()
        artifact = service.resume_job(job, worker_id="renderer-recovery")
        db.commit()

        assert artifact.id == job.artifact_id
        assert artifact.content_checksum == job.staged_checksum
