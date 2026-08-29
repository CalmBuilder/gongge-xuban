"""
@Time       : 2026/08/13 23:35
@Author     : zhanglp8181
@File       : artifact_renderer.py
@CallChain  : DynamicTaskAgent verified result → deterministic renderer → ArtifactService
@Description: 将已验证 Markdown 结果确定性渲染为文本、DOCX、CSV或XLSX并支持持久恢复。
"""

from __future__ import annotations

import hashlib
import csv
from io import BytesIO
from io import StringIO
from datetime import timedelta

from docx import Document
from openpyxl import Workbook
from sqlalchemy import update
from sqlmodel import Session, select

from app.db.models import (
    AgentProfile,
    ArtifactRendererJob,
    ChatSession,
    ExecutionArtifact,
    ExecutionResult,
    InputResourceSnapshot,
    SopInstance,
    SopNodeExecution,
    new_id,
    utc_now,
)
from app.dynamic_tasks.artifacts import ArtifactService


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
CSV_MIME = "text/csv"
RENDERER_VERSION = "deterministic-office-v1"
MAX_RENDER_ATTEMPTS = 3


class ArtifactRenderError(ValueError):
    """表示输出格式或输入内容违反确定性 renderer 契约。"""


class ArtifactRendererService:
    """以持久Job、lease和fencing驱动确定性产物生成及原子登记。"""

    def __init__(self, db: Session, *, artifact_service: ArtifactService | None = None) -> None:
        """绑定调用方事务和受管Artifact服务，测试可注入隔离存储。"""

        self.db = db
        self.artifacts = artifact_service or ArtifactService(db)

    def ensure_job(
        self,
        *,
        instance: SopInstance,
        result_id: str,
        result_checksum: str,
        source_node: SopNodeExecution,
        artifact_key: str,
        filename: str,
        mime_type: str,
        required: bool,
    ) -> tuple[ArtifactRendererJob, bool]:
        """幂等建立与已验证结果、节点和renderer版本绑定的渲染任务。"""

        existing = self.db.exec(
            select(ArtifactRendererJob).where(
                ArtifactRendererJob.tenant_id == instance.tenant_id,
                ArtifactRendererJob.execution_id == instance.id,
                ArtifactRendererJob.result_checksum == result_checksum,
                ArtifactRendererJob.artifact_key == artifact_key,
                ArtifactRendererJob.renderer_version == RENDERER_VERSION,
            )
        ).first()
        if existing is not None:
            if (
                existing.result_id != result_id
                or existing.source_node_execution_id != source_node.id
                or existing.filename != filename
                or existing.mime_type != mime_type
                or existing.required != required
            ):
                raise ArtifactRenderError("ARTIFACT_RENDER_JOB_IDENTITY_CONFLICT")
            return existing, False
        job = ArtifactRendererJob(
            tenant_id=instance.tenant_id,
            execution_id=instance.id,
            result_id=result_id,
            result_checksum=result_checksum,
            source_node_execution_id=source_node.id,
            artifact_key=artifact_key,
            filename=filename,
            mime_type=mime_type,
            renderer_version=RENDERER_VERSION,
            required=required,
            artifact_id=new_id("artifact"),
        )
        self.db.add(job)
        self.db.flush()
        return job, True

    def claim(
        self,
        job: ArtifactRendererJob,
        *,
        worker_id: str,
        lease_seconds: int = 300,
    ) -> ArtifactRendererJob:
        """以状态和到期租约CAS领取任务，旧worker只能得到过期fencing token。"""

        result = self.db.exec(
            update(ArtifactRendererJob)
            .where(
                ArtifactRendererJob.id == job.id,
                ArtifactRendererJob.status.in_(("pending", "retry_wait")),
                (
                    (ArtifactRendererJob.retry_at.is_(None))
                    | (ArtifactRendererJob.retry_at <= utc_now())
                ),
                ArtifactRendererJob.attempt_no < MAX_RENDER_ATTEMPTS,
            )
            .values(
                status="claimed",
                attempt_no=ArtifactRendererJob.attempt_no + 1,
                lease_owner=worker_id,
                lease_expires_at=utc_now() + timedelta(seconds=lease_seconds),
                fencing_token=ArtifactRendererJob.fencing_token + 1,
                updated_at=utc_now(),
            )
        )
        if result.rowcount != 1:
            raise ArtifactRenderError("ARTIFACT_RENDER_JOB_ALREADY_CLAIMED")
        self.db.flush()
        self.db.refresh(job)
        return job

    def requeue_expired(self) -> int:
        """回收过期渲染租约；耗尽尝试的任务直接死信，禁止形成永久重试等待。"""

        now = utc_now()
        expired_job_id = self.db.exec(
            select(ArtifactRendererJob.id)
            .where(
                ArtifactRendererJob.status.in_(("claimed", "rendering", "staged")),
                ArtifactRendererJob.lease_expires_at.is_not(None),
                ArtifactRendererJob.lease_expires_at <= now,
            )
            .limit(1)
        ).first()
        if expired_job_id is None:
            return 0
        terminal = self.db.exec(
            update(ArtifactRendererJob)
            .where(
                ArtifactRendererJob.status.in_(("claimed", "rendering", "staged")),
                ArtifactRendererJob.lease_expires_at.is_not(None),
                ArtifactRendererJob.lease_expires_at <= now,
                ArtifactRendererJob.attempt_no >= MAX_RENDER_ATTEMPTS,
            )
            .values(
                status="dead_letter",
                lease_owner=None,
                lease_expires_at=None,
                fencing_token=ArtifactRendererJob.fencing_token + 1,
                retry_at=None,
                error_code="ARTIFACT_RENDER_WORKER_LOST",
                updated_at=now,
            )
        )
        retryable = self.db.exec(
            update(ArtifactRendererJob)
            .where(
                ArtifactRendererJob.status.in_(("claimed", "rendering", "staged")),
                ArtifactRendererJob.lease_expires_at.is_not(None),
                ArtifactRendererJob.lease_expires_at <= now,
                ArtifactRendererJob.attempt_no < MAX_RENDER_ATTEMPTS,
            )
            .values(
                status="retry_wait",
                lease_owner=None,
                lease_expires_at=None,
                fencing_token=ArtifactRendererJob.fencing_token + 1,
                retry_at=now,
                error_code="ARTIFACT_RENDER_WORKER_LOST",
                updated_at=now,
            )
        )
        self.db.flush()
        return int(terminal.rowcount or 0) + int(retryable.rowcount or 0)

    def resume_job(self, job: ArtifactRendererJob, *, worker_id: str) -> ExecutionArtifact:
        """从持久ExecutionResult恢复渲染输入，重新校验血缘后完成一个已领取任务。"""

        self.db.refresh(job)
        result = self.db.get(ExecutionResult, job.result_id)
        if (
            result is None
            or result.tenant_id != job.tenant_id
            or result.execution_id != job.execution_id
            or result.status != "verified"
            or result.checksum != job.result_checksum
        ):
            raise ArtifactRenderError("ARTIFACT_RENDER_RESULT_INVALID")
        markdown = result.result_json.get("markdown") if isinstance(result.result_json, dict) else None
        if not isinstance(markdown, str) or not markdown.strip():
            raise ArtifactRenderError("ARTIFACT_RENDER_SOURCE_EMPTY")
        snapshot_ids = tuple(
            row.id
            for row in self.db.exec(
                select(InputResourceSnapshot).where(
                    InputResourceSnapshot.tenant_id == job.tenant_id,
                    InputResourceSnapshot.execution_id == job.execution_id,
                )
            ).all()
        )
        return self.render_and_publish(
            job,
            markdown=markdown,
            worker_id=worker_id,
            fencing_token=job.fencing_token,
            input_snapshot_ids=snapshot_ids,
        )

    def fail_or_retry(
        self,
        job: ArtifactRendererJob,
        *,
        worker_id: str,
        fencing_token: int,
        error_code: str,
    ) -> None:
        """以owner/fencing CAS收敛渲染故障，达到上限后进入可审计死信。"""

        terminal = job.attempt_no >= MAX_RENDER_ATTEMPTS
        result = self.db.exec(
            update(ArtifactRendererJob)
            .where(
                ArtifactRendererJob.id == job.id,
                ArtifactRendererJob.status.in_(("claimed", "rendering", "staged")),
                ArtifactRendererJob.lease_owner == worker_id,
                ArtifactRendererJob.fencing_token == fencing_token,
            )
            .values(
                status="dead_letter" if terminal else "retry_wait",
                lease_owner=None,
                lease_expires_at=None,
                retry_at=None if terminal else utc_now(),
                error_code=error_code[:64],
                updated_at=utc_now(),
            )
        )
        if result.rowcount != 1:
            raise ArtifactRenderError("ARTIFACT_RENDER_JOB_FENCED")
        self.db.flush()

    def render_and_publish(
        self,
        job: ArtifactRendererJob,
        *,
        markdown: str,
        worker_id: str,
        fencing_token: int,
        input_snapshot_ids: tuple[str, ...],
    ) -> ExecutionArtifact:
        """复核租约后渲染暂存内容，并以Artifact唯一键完成幂等发布。"""

        self.db.refresh(job)
        if (
            job.status != "claimed"
            or job.lease_owner != worker_id
            or job.fencing_token != fencing_token
            or job.lease_expires_at is None
            or job.lease_expires_at <= utc_now()
        ):
            raise ArtifactRenderError("ARTIFACT_RENDER_JOB_FENCED")
        instance = self.db.get(SopInstance, job.execution_id)
        source_node = self.db.get(SopNodeExecution, job.source_node_execution_id)
        result = self.db.get(ExecutionResult, job.result_id)
        if (
            instance is None
            or source_node is None
            or source_node.instance_id != instance.id
            or result is None
            or result.tenant_id != job.tenant_id
            or result.execution_id != instance.id
            or result.status != "verified"
            or result.checksum != job.result_checksum
            or result.result_json.get("markdown") != markdown
        ):
            raise ArtifactRenderError("ARTIFACT_RENDER_SOURCE_INVALID")
        self._lock_active_lineage_agent(instance)
        job.status = "rendering"
        self.db.add(job)
        self.db.flush()
        data = render_verified_markdown(markdown, job.mime_type)
        staged_checksum = hashlib.sha256(data).hexdigest()
        staged = self.db.exec(
            update(ArtifactRendererJob)
            .where(
                ArtifactRendererJob.id == job.id,
                ArtifactRendererJob.tenant_id == job.tenant_id,
                ArtifactRendererJob.status == "rendering",
                ArtifactRendererJob.lease_owner == worker_id,
                ArtifactRendererJob.fencing_token == fencing_token,
                ArtifactRendererJob.lease_expires_at > utc_now(),
            )
            .values(
                staged_checksum=staged_checksum,
                status="staged",
                updated_at=utc_now(),
            )
        )
        if getattr(staged, "rowcount", 0) != 1:
            self.db.rollback()
            raise ArtifactRenderError("ARTIFACT_RENDER_JOB_FENCED")
        self.db.refresh(job)
        artifact, _ = self.artifacts.register(
            instance=instance,
            source_node=source_node,
            artifact_key=job.artifact_key,
            filename=job.filename,
            mime_type=job.mime_type,
            data=data,
            input_snapshot_ids=input_snapshot_ids,
            artifact_id=job.artifact_id,
        )
        if artifact.content_checksum != job.staged_checksum:
            raise ArtifactRenderError("ARTIFACT_RENDER_STAGING_CHECKSUM_MISMATCH")
        ready = self.db.exec(
            update(ArtifactRendererJob)
            .where(
                ArtifactRendererJob.id == job.id,
                ArtifactRendererJob.tenant_id == job.tenant_id,
                ArtifactRendererJob.status == "staged",
                ArtifactRendererJob.lease_owner == worker_id,
                ArtifactRendererJob.fencing_token == fencing_token,
                ArtifactRendererJob.lease_expires_at > utc_now(),
            )
            .values(
                artifact_id=artifact.id,
                status="ready",
                lease_owner=None,
                lease_expires_at=None,
                updated_at=utc_now(),
            )
        )
        if getattr(ready, "rowcount", 0) != 1:
            self.db.rollback()
            raise ArtifactRenderError("ARTIFACT_RENDER_JOB_FENCED")
        self.db.refresh(job)
        return artifact

    def _lock_active_lineage_agent(self, instance: SopInstance) -> AgentProfile | None:
        """锁定 Execution 直接或会话继承的 Agent，渲染期间阻断墓碑竞态。"""

        agent_id = instance.agent_id
        if agent_id is None:
            session = self.db.get(ChatSession, instance.session_id)
            if session is not None and session.tenant_id == instance.tenant_id:
                agent_id = session.agent_id
        if agent_id is None:
            return None
        agent = self.db.exec(
            select(AgentProfile)
            .where(
                AgentProfile.tenant_id == instance.tenant_id,
                AgentProfile.id == agent_id,
            )
            .with_for_update()
            .execution_options(populate_existing=True)
        ).first()
        if agent is None or agent.status != "active" or agent.is_overall:
            raise ArtifactRenderError("ARTIFACT_RENDER_AGENT_DELETED")
        return agent


def render_verified_markdown(markdown: str, mime_type: str) -> bytes:
    """把已验证 Markdown 渲染为固定格式；不解释 HTML、公式、宏或外部链接。"""

    if not markdown.strip():
        raise ArtifactRenderError("ARTIFACT_RENDER_SOURCE_EMPTY")
    if mime_type == "text/markdown":
        return markdown.encode("utf-8")
    if mime_type == "text/plain":
        return markdown.encode("utf-8")
    if mime_type == CSV_MIME:
        return _render_csv(markdown)
    if mime_type == DOCX_MIME:
        return _render_docx(markdown)
    if mime_type == XLSX_MIME:
        return _render_xlsx(markdown)
    raise ArtifactRenderError("ARTIFACT_RENDER_MIME_UNSUPPORTED")


def safe_spreadsheet_text(value: str) -> str:
    """转义可能被表格软件解释为公式/DDE 的文本，同时保留用户可见内容。"""

    stripped = value.lstrip(" \t\r")
    if stripped.startswith(("=", "+", "-", "@")):
        return f"'{value}"
    return value


def _render_docx(markdown: str) -> bytes:
    """用 python-docx 生成无宏、无外链的基础报告。"""

    document = Document()
    for line in markdown.splitlines():
        text = line.strip()
        if not text:
            document.add_paragraph()
        elif text.startswith("### "):
            document.add_heading(text[4:], level=3)
        elif text.startswith("## "):
            document.add_heading(text[3:], level=2)
        elif text.startswith("# "):
            document.add_heading(text[2:], level=1)
        else:
            document.add_paragraph(text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _render_xlsx(markdown: str) -> bytes:
    """把报告逐行写入只含文本的工作簿，避免模型文本触发表格公式执行。"""

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "报告"
    sheet.append(["序号", "内容"])
    for index, line in enumerate(markdown.splitlines(), start=1):
        sheet.append([index, safe_spreadsheet_text(line)])
    sheet.freeze_panes = "A2"
    sheet.column_dimensions["A"].width = 10
    sheet.column_dimensions["B"].width = 100
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _render_csv(markdown: str) -> bytes:
    """把报告逐行渲染为可重读UTF-8 CSV，并阻断公式与DDE单元格注入。"""

    output = StringIO(newline="")
    writer = csv.writer(output, lineterminator="\r\n")
    writer.writerow(["序号", "内容"])
    for index, line in enumerate(markdown.splitlines(), start=1):
        writer.writerow([index, safe_spreadsheet_text(line)])
    return output.getvalue().encode("utf-8-sig")
